import logging
import os
import re
import zipfile
from io import BytesIO

from docx import Document
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.models import Conversation, Contract, ContractTemplate
from app.services.llm_service import generate_contract

logger = logging.getLogger(__name__)


def _document_to_plain_text(doc: Document) -> str:
    parts: list[str] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts) if parts else ""


def extract_docx_text(file_path: str) -> str:
    """Read plain text from a .docx file (paragraphs and table cells)."""
    return _document_to_plain_text(Document(file_path))


def extract_docx_text_from_bytes(data: bytes) -> str:
    """Read plain text from .docx bytes (e.g. multipart upload without saving to disk)."""
    return _document_to_plain_text(Document(BytesIO(data)))


def _split_contract_title_and_body(raw: str) -> tuple[str, str]:
    """Parse LLM output: TITLE: <short>\\n---\\n<body>. Fallback: first line as title, rest as body."""
    text = raw.strip()
    if not text:
        return "合同草稿", ""

    m = re.match(r"(?is)^TITLE:\s*([^\n]+?)\s*\n\s*---+?\s*\n(.*)$", text)
    if m:
        title = m.group(1).strip().strip('"').strip("'")
        body = m.group(2).strip()
        if title:
            if len(title) > 200:
                title = title[:80].rstrip() + "…"
            return title[:500], body if body else text

    m2 = re.match(r"(?is)^TITLE:\s*([^\n]+)\s*\n+(.*)$", text)
    if m2:
        title = m2.group(1).strip().strip('"').strip("'")
        body = m2.group(2).strip()
        if title and len(body) > 20:
            if len(title) > 200:
                title = title[:80].rstrip() + "…"
            return title[:500], body

    lines = text.split("\n")
    first = lines[0].strip().strip("#").strip("*").strip()
    if first.lower().startswith("title:"):
        first = first[6:].strip()
    title = first[:500] if len(first) <= 120 else (first[:72].rstrip() + "…")
    rest = "\n".join(lines[1:]).strip()
    return title, rest if rest else text


def looks_like_ooxml_docx(data: bytes) -> bool:
    """True if bytes look like a .docx (ZIP OOXML with word/document.xml). Catches uploads with missing/wrong filename."""
    if len(data) < 4 or data[:2] != b"PK":
        return False
    try:
        with zipfile.ZipFile(BytesIO(data)) as zf:
            return "word/document.xml" in zf.namelist()
    except zipfile.BadZipFile:
        return False


async def create_contract_from_conversation(
    db: AsyncSession,
    conversation_id: int,
    template_id: int | None = None,
    output_language: str | None = None,
) -> Contract | None:
    """Generate a contract draft from a conversation's chat history."""
    result = await db.execute(
        select(Conversation)
        .options(selectinload(Conversation.messages))
        .where(Conversation.id == conversation_id)
    )
    conversation = result.scalar_one_or_none()
    if not conversation:
        return None

    chat_history = [
        {"role": msg.role.value, "content": msg.content}
        for msg in conversation.messages
    ]

    customer_name = " ".join(
        filter(None, [conversation.first_name, conversation.last_name])
    ) or conversation.username or f"Customer #{conversation.telegram_user_id}"

    template_content: str | None = None
    if template_id is not None:
        template = await db.get(ContractTemplate, template_id)
        if template and template.file_path and os.path.exists(template.file_path):
            try:
                template_content = extract_docx_text(template.file_path)
            except Exception as e:
                logger.error(f"Failed to read template {template_id}: {e}")
                template_content = None
        elif template_id is not None:
            logger.warning(f"Template {template_id} not found or file missing")

    lang = (output_language or "").strip() or conversation.language or "en"

    raw = await generate_contract(
        chat_history=chat_history,
        customer_name=customer_name,
        language=lang,
        template_content=template_content,
    )

    title_line, body = _split_contract_title_and_body(raw)
    contract = Contract(
        conversation_id=conversation_id,
        title=title_line[:500],
        content=body,
        status="draft",
    )
    db.add(contract)
    await db.commit()
    await db.refresh(contract)
    return contract
